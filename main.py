import os
import re
import uuid
import time
import base64
import urllib.parse
import datetime
import textwrap
import traceback
from io import BytesIO
from pathlib import Path
from docx import Document
import json
from fastapi.responses import StreamingResponse

import requests
import edge_tts

from dotenv import load_dotenv
from bson import ObjectId

from fastapi import FastAPI, HTTPException, Body, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, FileResponse
from fastapi.staticfiles import StaticFiles

from google import genai
from google.genai import types
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests

from pypdf import PdfReader

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

from PyPDF2 import PdfReader
from docx import Document

from database import (
    users_collection,
    chat_collection,
    conversations_collection,
    images_collection,
    audio_collection,
    notes_collection,
    documents_collection,
    credit_logs_collection
)

from auth import hash_password, verify_password, create_token
from datetime import datetime, timezone

from models import (
    UserSignup,
    UserLogin,
    ChatRequest,
    UserProfileUpdate
)

import secrets
from datetime import datetime, timedelta
from fastapi.responses import RedirectResponse

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart


# =========================
# SETUP
# =========================

load_dotenv(dotenv_path=".env")

FRONTEND_URL = os.getenv("FRONTEND_URL", "http://127.0.0.1:5500/project")
BACKEND_URL = os.getenv("BACKEND_URL", "http://127.0.0.1:8000")

SMTP_EMAIL = os.getenv("SMTP_EMAIL")
SMTP_APP_PASSWORD = os.getenv("SMTP_APP_PASSWORD")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "models/gemma-4-26b-a4b-it")

GENERATED_IMAGE_DIR = "generated_images"
GENERATED_AUDIO_DIR = "generated_audio"
EXPORT_DIR = "exports"

os.makedirs(GENERATED_IMAGE_DIR, exist_ok=True)
os.makedirs(GENERATED_AUDIO_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

gemini_client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

app = FastAPI(title="Lumina AI Backend")
app.mount("/static", StaticFiles(directory="."), name="static")

UPLOAD_DOCUMENTS_DIR = Path.home() / "lumina_storage" / "uploaded_documents"
UPLOAD_DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

app.mount(
    "/uploaded_documents",
    StaticFiles(directory=str(UPLOAD_DOCUMENTS_DIR)),
    name="uploaded_documents"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount(
    "/generated_images",
    StaticFiles(directory=GENERATED_IMAGE_DIR),
    name="generated_images"
)

app.mount(
    "/generated_audio",
    StaticFiles(directory=GENERATED_AUDIO_DIR),
    name="generated_audio"
)

LUMINA_PLANS = {
    "free": {
        "name": "Free",
        "price": 0,
        "monthly_credits": 3000,
        "daily_limit": 300
    },
    "starter": {
        "name": "Starter",
        "price": 99,
        "monthly_credits": 10000,
        "daily_limit": 700
    },
    "creator": {
        "name": "Creator",
        "price": 199,
        "monthly_credits": 25000,
        "daily_limit": 1500
    },
    "pro": {
        "name": "Pro",
        "price": 499,
        "monthly_credits": 75000,
        "daily_limit": 5000
    }
}

CREDIT_COSTS = {
    "chat_message": 1,
    "chat_file_upload": 5,
    "document_analysis": 8,
    "image_generation": 15,
    "voice_generation": 10
}

# =========================
# HELPERS
# =========================

def send_verification_email(to_email: str, verification_link: str):
    if not SMTP_EMAIL or not SMTP_APP_PASSWORD:
        raise Exception("SMTP_EMAIL or SMTP_APP_PASSWORD missing in .env")

    subject = "Verify your Lumina account"

    html_body = f"""
    <div style="font-family: Arial, sans-serif; padding: 24px; background: #f8fbff;">
      <div style="max-width: 560px; margin: auto; background: white; padding: 28px; border-radius: 18px;">
        <h2 style="color: #2563eb;">Verify your Lumina account</h2>

        <p style="color: #334155; line-height: 1.6;">
          Thanks for creating your Lumina account. Please verify your Gmail address before purchasing a Lumina plan.
        </p>

        <a href="{verification_link}"
           style="display: inline-block; margin-top: 16px; padding: 14px 22px; background: #2563eb; color: white; text-decoration: none; border-radius: 999px; font-weight: bold;">
          Verify Email
        </a>

        <p style="margin-top: 22px; color: #64748b; font-size: 13px;">
          We will never ask for your OTP. You only need to click this verification link.
        </p>
      </div>
    </div>
    """

    message = MIMEMultipart("alternative")
    message["Subject"] = subject
    message["From"] = SMTP_EMAIL
    message["To"] = to_email
    message.attach(MIMEText(html_body, "html"))

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(SMTP_EMAIL, SMTP_APP_PASSWORD)
        server.sendmail(SMTP_EMAIL, to_email, message.as_string())

def current_month_key():
    return datetime.now(timezone.utc).strftime("%Y-%m")


def current_date_key():
    return datetime.now(timezone.utc).strftime("%Y-%m-%d")


def get_plan_details(plan_name: str):
    plan_name = str(plan_name or "free").lower()
    return LUMINA_PLANS.get(plan_name, LUMINA_PLANS["free"])


def reset_monthly_credits_if_needed(user_id: str):
    if str(user_id).startswith("guest_"):
        return

    user = users_collection.find_one({"_id": object_id(user_id)})

    if not user:
        return

    current_month = current_month_key()
    last_reset_month = user.get("last_credit_reset_month")

    if last_reset_month == current_month:
        return

    plan = user.get("plan", "free")
    plan_details = get_plan_details(plan)

    users_collection.update_one(
        {"_id": object_id(user_id)},
        {
            "$set": {
                "credits": plan_details["monthly_credits"],
                "monthly_credit_limit": plan_details["monthly_credits"],
                "daily_credit_limit": plan_details["daily_limit"],
                "daily_used_credits": 0,
                "last_credit_reset_month": current_month,
                "last_daily_reset_date": current_date_key(),
                "plan": plan
            }
        }
    )

def reset_daily_credits_if_needed(user_id: str):
    if str(user_id).startswith("guest_"):
        return

    user = users_collection.find_one({"_id": object_id(user_id)})

    if not user:
        return

    today = current_date_key()
    last_daily_reset = user.get("last_daily_reset_date")

    if last_daily_reset == today:
        return

    plan = user.get("plan", "free")
    plan_details = get_plan_details(plan)

    users_collection.update_one(
        {"_id": object_id(user_id)},
        {
            "$set": {
                "daily_used_credits": 0,
                "daily_credit_limit": plan_details["daily_limit"],
                "last_daily_reset_date": today
            }
        }
    )

def get_user_credits(user_id: str) -> int:
    if str(user_id).startswith("guest_"):
        return 0

    reset_monthly_credits_if_needed(user_id)

    user = users_collection.find_one({"_id": object_id(user_id)})

    if not user:
        return 0

    return int(user.get("credits", 0))


def deduct_user_credits(user_id: str, cost: int, reason: str):
    if str(user_id).startswith("guest_"):
        raise HTTPException(
            status_code=403,
            detail="Please login to use this feature."
        )

    reset_monthly_credits_if_needed(user_id)
    reset_daily_credits_if_needed(user_id)

    user = users_collection.find_one({"_id": object_id(user_id)})

    if not user:
        raise HTTPException(status_code=404, detail="User not found.")

    current_credits = int(user.get("credits", 0))
    daily_used = int(user.get("daily_used_credits", 0))
    daily_limit = int(user.get("daily_credit_limit", 300))

    if current_credits < cost:
        raise HTTPException(
            status_code=402,
            detail="Not enough Lumina Credits."
        )

    if daily_used + cost > daily_limit:
        raise HTTPException(
            status_code=429,
            detail=f"Daily credit limit reached. You can use up to {daily_limit} credits per day on your current plan."
        )

    users_collection.update_one(
        {"_id": object_id(user_id)},
        {
            "$inc": {
                "credits": -cost,
                "daily_used_credits": cost
            }
        }
    )

    try:
        credit_logs_collection.insert_one({
            "user_id": user_id,
            "type": "deduct",
            "cost": cost,
            "reason": reason,
            "daily_used_after": daily_used + cost,
            "daily_limit": daily_limit,
            "created_at": now_iso()
        })
    except Exception as e:
        print("CREDIT LOG ERROR:", e)


def refund_user_credits(user_id: str, amount: int, reason: str):
    if str(user_id).startswith("guest_"):
        return

    users_collection.update_one(
        {"_id": object_id(user_id)},
        {
            "$inc": {
                "credits": amount,
                "daily_used_credits": -amount
            }
        }
    )

    users_collection.update_one(
        {
            "_id": object_id(user_id),
            "daily_used_credits": {"$lt": 0}
        },
        {
            "$set": {
                "daily_used_credits": 0
            }
        }
    )

    try:
        credit_logs_collection.insert_one({
            "user_id": user_id,
            "type": "refund",
            "amount": amount,
            "reason": reason,
            "created_at": now_iso()
        })
    except Exception as e:
        print("CREDIT REFUND LOG ERROR:", e)

def now_iso():
    return datetime.now(timezone.utc).isoformat()


def object_id(id_value: str):
    try:
        return ObjectId(id_value)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid ID")


def serialize_doc(doc):
    if not doc:
        return None

    doc["_id"] = str(doc["_id"])
    return doc


def serialize_docs(docs):
    return [serialize_doc(doc) for doc in docs]


def is_strong_password(password: str):
    if len(password) < 8:
        return False

    if not re.search(r"[A-Z]", password):
        return False

    if not re.search(r"[0-9]", password):
        return False

    if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):
        return False

    return True


def ensure_gemini_ready():
    if not GEMINI_API_KEY or not gemini_client:
        raise HTTPException(
            status_code=500,
            detail="GEMINI_API_KEY is missing in .env"
        )


def get_size_from_ratio(aspect_ratio: str):
    ratio_map = {
        "Square": "1024x1024",
        "Portrait": "1024x1536",
        "Landscape": "1536x1024",
        "1:1": "1024x1024",
        "4:5": "1024x1536",
        "16:9": "1536x1024"
    }

    return ratio_map.get(aspect_ratio, "1024x1024")


def build_style_prompt(user_prompt: str, style: str, has_reference: bool = False):
    style_map = {
        "Realistic": "highly detailed, photorealistic, sharp focus, realistic lighting",
        "Anime": "anime style, clean line art, vibrant colors, expressive, polished illustration",
        "Cyberpunk": "cyberpunk aesthetic, neon lights, futuristic mood, cinematic atmosphere",
        "3D": "3D rendered, detailed materials, realistic shading, polished lighting",
        "Logo": "clean logo design, minimalist, vector-like, centered composition",
        "Watercolor": "watercolor painting, soft edges, brush texture, artistic and elegant"
    }

    style_text = style_map.get(style, "high quality, detailed")

    if has_reference:
        return f"""
Use the uploaded image as the primary reference.

User request:
{user_prompt}

Style:
{style_text}

Instructions:
- Preserve the main subject, identity, pose, and overall composition unless the user explicitly asks to change them.
- Keep the result visually faithful to the uploaded image.
- Improve quality, lighting, and style.
- Do not add unrelated objects or change the scene completely.
- Generate a polished final image.
"""

    return f"""
Create an image based on this request:

{user_prompt}

Style:
{style_text}

Instructions:
- Make it visually strong, high quality, sharp, and detailed.
- Follow the prompt closely.
- Keep the composition clean and attractive.
"""

def generate_ai_response(prompt: str) -> str:
    try:
        if not prompt or not prompt.strip():
            return "No prompt provided."

        response = GEMINI_MODEL(prompt)

        if hasattr(response, "text") and response.text:
            return response.text

        return "Lumina could not generate a response."

    except Exception as e:
        print("AI RESPONSE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"AI response failed: {str(e)}"
        )

# =========================
# ROOT / HEALTH
# =========================

from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles

@app.get("/")
def home():
    return FileResponse("landing.html")

@app.get("/debug-routes")
def debug_routes():
    return [
        {
            "path": route.path,
            "methods": list(route.methods)
        }
        for route in app.routes
        if hasattr(route, "methods")
    ]
# =========================
# AUTH
# =========================

@app.post("/signup")
def signup(user: UserSignup):
    try:
        if not is_strong_password(user.password):
            raise HTTPException(
                status_code=400,
                detail="Password must be at least 8 characters and include 1 uppercase letter, 1 number, and 1 symbol."
            )

        email = user.email.strip().lower()
        username = user.username.strip()

        if users_collection.find_one({"email": email}):
            raise HTTPException(
                status_code=400,
                detail="Email already exists"
            )

        if users_collection.find_one({"username": username}):
            suggestions = [
                username + "123",
                username + "_ai",
                username + "_lumina"
            ]

            raise HTTPException(
                status_code=400,
                detail={
                    "message": "Username already exists",
                    "suggestions": suggestions
                }
            )

        verification_token = secrets.token_urlsafe(32)
        verification_expires = datetime.utcnow() + timedelta(hours=24)
        verification_link = f"{BACKEND_URL}/verify-email?token={verification_token}"

        new_user = {
            "full_name": user.full_name,
            "email": email,
            "username": username,
            "password": hash_password(user.password),
            "auth_provider": "email",

            "phone": "",
            "bio": "",
            "occupation": "",
            "profile_photo": "",

            # Gmail verification
            "email_verified": False,
            "email_verification_token": verification_token,
            "email_verification_expires": verification_expires,

            # Plan / credits
            "plan": "free",
            "credits": 3000,
            "monthly_credit_limit": 3000,
            "daily_credit_limit": 300,
            "daily_used_credits": 0,
            "last_credit_reset_month": current_month_key(),
            "last_daily_reset_date": current_date_key(),

            "created_at": now_iso()
        }

        result = users_collection.insert_one(new_user)

        verification_email_sent = False

        try:
            send_verification_email(email, verification_link)
            verification_email_sent = True

        except Exception as email_error:
            print("VERIFICATION EMAIL ERROR:", email_error)
            traceback.print_exc()
            verification_email_sent = False

        token = create_token({
            "email": email
        })

        return {
            "message": "User created successfully",
            "verification_message": (
                "Verification link sent to your Gmail."
                if verification_email_sent
                else "Account created, but verification email could not be sent. Please resend verification from Plans page."
            ),
            "verification_email_sent": verification_email_sent,
            "token": token,
            "user_id": str(result.inserted_id),
            "full_name": user.full_name,
            "email": email,
            "username": username,
            "profile_photo": "",
            "email_verified": False,
            "plan": "free",
            "credits": 3000,
            "monthly_credit_limit": 3000,
            "daily_credit_limit": 300
        }

    except HTTPException:
        raise

    except Exception as e:
        print("SIGNUP ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="Signup failed"
        )
    
@app.post("/login")
def login(user: UserLogin):
    try:
        db_user = users_collection.find_one({"email": user.email})

        if not db_user:
            raise HTTPException(status_code=400, detail="Invalid credentials")

        if db_user.get("auth_provider") == "google":
            raise HTTPException(status_code=400, detail="This account uses Google login")

        if not verify_password(user.password, db_user.get("password")):
            raise HTTPException(status_code=400, detail="Invalid credentials")

        token = create_token({"email": user.email})

        return {
            "message": "Login successful",
            "token": token,
            "user_id": str(db_user["_id"]),
            "full_name": db_user.get("full_name", ""),
            "email": db_user.get("email", ""),
            "username": db_user.get("username", ""),
            "profile_photo": db_user.get("profile_photo", "")
        }

    except HTTPException:
        raise

    except Exception as e:
        print("LOGIN ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Login failed")


@app.post("/google-login")
def google_login(data: dict):
    try:
        token = data.get("credential")

        if not token:
            raise HTTPException(status_code=400, detail="Google token missing")

        if not GOOGLE_CLIENT_ID:
            raise HTTPException(status_code=500, detail="Google Client ID missing in .env")

        id_info = id_token.verify_oauth2_token(
            token,
            google_requests.Request(),
            GOOGLE_CLIENT_ID
        )

        email = (id_info.get("email") or "").strip().lower()
        full_name = id_info.get("name") or "Google User"
        picture = id_info.get("picture") or ""
        google_id = id_info.get("sub")

        if not email:
            raise HTTPException(status_code=400, detail="Google email not found")

        user = users_collection.find_one({"email": email})

        if not user:
            username_base = email.split("@")[0]

            result = users_collection.insert_one({
                "full_name": full_name,
                "email": email,
                "username": username_base,
                "password": None,
                "auth_provider": "google",
                "google_id": google_id,
                "profile_photo": picture,
                "phone": "",
                "bio": "",
                "occupation": "",

                "email_verified": True,
                "email_verification_token": "",
                "email_verification_expires": None,

                "plan": "free",
                "credits": 3000,
                "monthly_credit_limit": 3000,
                "daily_credit_limit": 300,
                "daily_used_credits": 0,
                "last_credit_reset_month": current_month_key(),
                "last_daily_reset_date": current_date_key(),

                "created_at": now_iso()
            })

            user_id = str(result.inserted_id)
            username = username_base

        else:
            user_id = str(user["_id"])
            username = user.get("username", email.split("@")[0])
            full_name = user.get("full_name", full_name)
            picture = user.get("profile_photo", picture)

            users_collection.update_one(
                {"_id": user["_id"]},
                {
                    "$set": {
                        "google_id": google_id,
                        "auth_provider": "google",
                        "email_verified": True,
                        "email_verification_token": "",
                        "email_verification_expires": None
                    }
                }
            )

        app_token = create_token({"email": email})

        return {
            "message": "Google login successful",
            "token": app_token,
            "user_id": user_id,
            "full_name": full_name,
            "email": email,
            "username": username,
            "profile_photo": picture,
            "email_verified": True
        }

    except ValueError as e:
        print("GOOGLE TOKEN VERIFY ERROR:", e)
        raise HTTPException(status_code=401, detail=f"Invalid Google token: {str(e)}")

    except HTTPException:
        raise

    except Exception as e:
        print("GOOGLE LOGIN ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Google login failed")

@app.get("/verify-email")
def verify_email(token: str):
    try:
        user = users_collection.find_one({
            "email_verification_token": token
        })

        if not user:
            return RedirectResponse(
                url=f"{FRONTEND_URL}/verify-success.html?status=invalid"
            )

        expires = user.get("email_verification_expires")

        if isinstance(expires, str):
            try:
                expires = datetime.fromisoformat(expires.replace("Z", "+00:00"))
            except Exception:
                expires = None

        if expires:
            now = datetime.utcnow()

            if getattr(expires, "tzinfo", None) is not None:
                expires = expires.replace(tzinfo=None)

            if now > expires:
                return RedirectResponse(
                    url=f"{FRONTEND_URL}/verify-success.html?status=expired"
                )

        users_collection.update_one(
            {"_id": user["_id"]},
            {
                "$set": {
                    "email_verified": True,
                    "email_verification_token": "",
                    "email_verification_expires": None
                }
            }
        )

        return RedirectResponse(
            url=f"{FRONTEND_URL}/verify-success.html?status=success"
        )

    except Exception as e:
        print("VERIFY EMAIL ERROR:", e)
        traceback.print_exc()
        return RedirectResponse(
            url=f"{FRONTEND_URL}/verify-success.html?status=invalid"
        )
        
@app.post("/resend-verification")
def resend_verification(data: dict):
    try:
        email = str(data.get("email", "")).strip().lower()

        if not email:
            raise HTTPException(status_code=400, detail="Email is required")

        user = users_collection.find_one({"email": email})

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        if user.get("email_verified", False):
            return {
                "success": True,
                "already_verified": True,
                "message": "Email is already verified"
            }

        verification_token = secrets.token_urlsafe(32)
        verification_expires = datetime.utcnow() + timedelta(hours=24)
        verification_link = f"{BACKEND_URL}/verify-email?token={verification_token}"

        users_collection.update_one(
            {"email": email},
            {
                "$set": {
                    "email_verified": False,
                    "email_verification_token": verification_token,
                    "email_verification_expires": verification_expires
                }
            }
        )

        send_verification_email(email, verification_link)

        return {
            "success": True,
            "already_verified": False,
            "message": "Verification email sent"
        }

    except HTTPException:
        raise

    except Exception as e:
        print("RESEND VERIFICATION ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="Could not send verification email."
        )
    
# =========================
# PROFILE
# =========================

@app.get("/profile/{user_id}")
def get_profile(user_id: str):
    try:
        user = users_collection.find_one({"_id": object_id(user_id)})

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        return {
            "user_id": str(user["_id"]),
            "full_name": user.get("full_name", "User"),
            "email": user.get("email", ""),
            "username": user.get("username", ""),
            "phone": user.get("phone", ""),
            "bio": user.get("bio", ""),
            "occupation": user.get("occupation", ""),
            "profile_photo": user.get("profile_photo", ""),
            "auth_provider": user.get("auth_provider", "email"),
            "email_verified": user.get("email_verified", False),
            "plan": user.get("plan", "free"),
            "credits": int(user.get("credits", 0)),
            "monthly_credit_limit": int(user.get("monthly_credit_limit", 3000)),
            "daily_credit_limit": int(user.get("daily_credit_limit", 300)),
            "daily_used_credits": int(user.get("daily_used_credits", 0))
        }

    except HTTPException:
        raise

    except Exception as e:
        print("GET PROFILE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load profile")
    
@app.get("/user-profile/{user_id}")
def get_user_profile_alias(user_id: str):
    return get_profile(user_id)

@app.put("/profile/{user_id}")
def update_profile(user_id: str, data: UserProfileUpdate):
    try:
        update_data = {
            "full_name": data.full_name or "",
            "phone": data.phone or "",
            "bio": data.bio or "",
            "occupation": data.occupation or ""
        }

        result = users_collection.update_one(
            {"_id": object_id(user_id)},
            {"$set": update_data}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="User not found")

        user = users_collection.find_one({"_id": object_id(user_id)})

        return {
            "message": "Profile updated successfully",
            "user_id": str(user["_id"]),
            "full_name": user.get("full_name", "User"),
            "email": user.get("email", ""),
            "username": user.get("username", ""),
            "phone": user.get("phone", ""),
            "bio": user.get("bio", ""),
            "occupation": user.get("occupation", ""),
            "profile_photo": user.get("profile_photo", ""),
            "auth_provider": user.get("auth_provider", "email")
        }

    except HTTPException:
        raise

    except Exception as e:
        print("UPDATE PROFILE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not update profile")


@app.put("/update-account/{user_id}")
def update_account(user_id: str, data: dict):
    try:
        full_name = data.get("full_name", "").strip()
        username = data.get("username", "").strip()

        if not full_name or not username:
            raise HTTPException(status_code=400, detail="Name and username are required")

        existing_user = users_collection.find_one({
            "username": username,
            "_id": {"$ne": object_id(user_id)}
        })

        if existing_user:
            raise HTTPException(status_code=400, detail="Username already exists")

        result = users_collection.update_one(
            {"_id": object_id(user_id)},
            {"$set": {
                "full_name": full_name,
                "username": username
            }}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="User not found")

        return {
            "message": "Account updated",
            "full_name": full_name,
            "username": username
        }

    except HTTPException:
        raise

    except Exception as e:
        print("UPDATE ACCOUNT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not update account")


@app.put("/change-password")
def change_password(data: dict):
    try:
        user_id = data.get("user_id")
        old_password = data.get("old_password")
        new_password = data.get("new_password")

        if not user_id or not old_password or not new_password:
            raise HTTPException(status_code=400, detail="All fields are required")

        if not is_strong_password(new_password):
            raise HTTPException(
                status_code=400,
                detail="New password must be at least 8 characters and include 1 uppercase letter, 1 number, and 1 symbol."
            )

        user = users_collection.find_one({"_id": object_id(user_id)})

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        if user.get("auth_provider") == "google":
            raise HTTPException(status_code=400, detail="Google accounts cannot change password here")

        if not verify_password(old_password, user.get("password")):
            raise HTTPException(status_code=400, detail="Old password is incorrect")

        users_collection.update_one(
            {"_id": object_id(user_id)},
            {"$set": {"password": hash_password(new_password)}}
        )

        return {"message": "Password changed successfully"}

    except HTTPException:
        raise

    except Exception as e:
        print("CHANGE PASSWORD ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not change password")


@app.delete("/delete-account")
def delete_account(data: dict = Body(...)):
    try:
        user_id = data.get("user_id")
        password = data.get("password")

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        user = users_collection.find_one({"_id": object_id(user_id)})

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        if user.get("auth_provider") != "google":
            if not password:
                raise HTTPException(status_code=400, detail="Password is required")

            if not verify_password(password, user.get("password")):
                raise HTTPException(status_code=400, detail="Incorrect password")

        user_images = list(images_collection.find({"user_id": str(user_id)}))

        for image in user_images:
            image_url = image.get("image_url", "")

            if image_url.startswith("/generated_images/"):
                file_name = image_url.replace("/generated_images/", "")
                file_path = os.path.join(GENERATED_IMAGE_DIR, file_name)

                if os.path.exists(file_path):
                    os.remove(file_path)

        user_audio = list(audio_collection.find({"user_id": str(user_id)}))

        for audio in user_audio:
            audio_url = audio.get("audio_url", "")

            if audio_url.startswith("/generated_audio/"):
                file_name = audio_url.replace("/generated_audio/", "")
                file_path = os.path.join(GENERATED_AUDIO_DIR, file_name)

                if os.path.exists(file_path):
                    os.remove(file_path)

        users_collection.delete_one({"_id": object_id(user_id)})
        conversations_collection.delete_many({"user_id": str(user_id)})
        chat_collection.delete_many({"user_id": str(user_id)})
        images_collection.delete_many({"user_id": str(user_id)})
        audio_collection.delete_many({"user_id": str(user_id)})
        notes_collection.delete_many({"user_id": str(user_id)})

        return {"message": "Account deleted"}

    except HTTPException:
        raise

    except Exception as e:
        print("DELETE ACCOUNT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete account")


# =========================
# CHAT
# =========================

@app.post("/new-chat")
def new_chat(data: dict):
    try:
        user_id = str(data.get("user_id", "")).strip()

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        chat = {
            "user_id": user_id,
            "title": "New Chat",
            "created_at": now_iso(),
            "pinned": False,
            "archived": False
        }

        result = conversations_collection.insert_one(chat)

        return {"chat_id": str(result.inserted_id)}

    except HTTPException:
        raise

    except Exception as e:
        print("NEW CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not create chat")


def get_ai_response(message: str, history_text: str = ""):
    try:
        if not GEMINI_API_KEY or not gemini_client:
            print("GEMINI ERROR: API key missing from .env")
            return "Gemini API key is missing. Please check your .env file."

        final_prompt = f"""
You are Lumina AI, a helpful, intelligent, and well-structured AI assistant.

Your goal:
Give answers that are clear, organized, useful, and visually structured.

GENERAL ANSWER RULES:
- Use clean Markdown formatting.
- Use headings, subheadings, bullet points, numbered steps, and tables when useful.
- Do not write messy long paragraphs.
- Keep the answer easy to scan.
- Use simple language unless the user asks for advanced detail.
- If the user asks for comparison, use a Markdown table.
- If the user asks for steps, use numbered steps.
- If the user asks for process, lifecycle, architecture, workflow, or decision logic, include a Mermaid flowchart.
- If the user asks for code, provide clean code blocks with language labels.
- If the user asks for exam/viva answers, give clear academic-style answers.
- If the answer is long, divide it into sections.
- End with a short summary when useful.

FLOWCHART RULES:
When explaining a process, workflow, lifecycle, architecture, or step-by-step system, include Mermaid like this:

```mermaid
flowchart TD
    A[Start] --> B[Step 1]
    B --> C[Step 2]
    C --> D[End]
```

IMAGE PROMPT RULES:
If user asks for an image, return:

## Image Prompt
[Detailed prompt]

## Style Suggestions
- Realistic
- Anime
- 3D
- Cyberpunk

## Aspect Ratio
Recommended: Square / Portrait / Landscape

Previous conversation:
{history_text}

User:
{message}

Lumina:
"""

        models_to_try = [
    "gemini-2.5-flash",
    "gemini-1.5-flash",
    "models/gemma-4-26b-a4b-it"
]

        # remove duplicates while keeping order
        unique_models = []
        for model in models_to_try:
            if model and model not in unique_models:
                unique_models.append(model)

        last_error = None

        for model_name in unique_models:
            try:
                print("TRYING GEMINI MODEL:", model_name)

                response = gemini_client.models.generate_content(
                    model=model_name,
                    contents=final_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.7,
                        max_output_tokens=1200
                    )
                )

                if response and getattr(response, "text", None):
                    print("WORKING MODEL:", model_name)
                    return response.text.strip()

            except Exception as e:
                last_error = e
                print(f"GEMINI ERROR WITH {model_name}:", e)

        print("ALL GEMINI MODELS FAILED:", last_error)
        return "Lumina could not connect to Gemini right now. Please try again."

    except Exception as e:
        print("AI FUNCTION ERROR:", e)
        traceback.print_exc()
        return "Lumina is temporarily unavailable. Please try again."

def stream_ai_response(message: str, history_text: str = ""):
    if not GEMINI_API_KEY or not gemini_client:
        yield "Gemini API key is missing. Please check your .env file."
        return

    final_prompt = f"""
You are Lumina AI, a helpful and well-structured AI assistant.

Answer rules:
- Use clean Markdown.
- Use headings, bullets, steps, and tables when useful.
- Keep answers clear and easy to scan.
- Use simple language unless advanced detail is requested.
- For workflows/processes, include Mermaid flowchart only if useful.
- For code, use clean code blocks.
- For image requests, give an Image Prompt, Style Suggestions, and Aspect Ratio.

Previous conversation:
{history_text}

User:
{message}

Lumina:
"""

    models_to_try = [
    "gemini-2.5-flash",
    "gemini-1.5-flash",
    "models/gemma-4-26b-a4b-it"
]

    unique_models = []

    for model in models_to_try:
        if model and model not in unique_models:
            unique_models.append(model)

    last_error = None

    for model_name in unique_models:
        try:
            print("TRYING STREAMING GEMINI MODEL:", model_name)

            stream = gemini_client.models.generate_content_stream(
    model=model_name,
    contents=final_prompt,
    config=types.GenerateContentConfig(
        temperature=0.5,
        max_output_tokens=1200
    )
)

            got_text = False

            for chunk in stream:
                chunk_text = getattr(chunk, "text", None)

                if chunk_text:
                    got_text = True
                    yield chunk_text

            if got_text:
                print("WORKING STREAMING MODEL:", model_name)
                return

        except Exception as e:
            last_error = e
            print(f"STREAMING GEMINI ERROR WITH {model_name}:", e)

    print("ALL STREAMING GEMINI MODELS FAILED:", last_error)
    yield "Lumina could not connect to Gemini right now. Please try again."

@app.post("/chat")
def chat(data: ChatRequest):
    try:
        user_id = str(data.user_id).strip()
        chat_id = str(data.chat_id).strip()
        message = data.message.strip()

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        if not chat_id:
            raise HTTPException(status_code=400, detail="Chat ID is required")

        if not message:
            raise HTTPException(status_code=400, detail="Message is required")

        previous_messages = list(
            chat_collection
            .find({"chat_id": chat_id})
            .sort("_id", -1)
            .limit(6)
        )

        previous_messages.reverse()

        history_text = ""

        for msg in previous_messages:
            history_text += f"User: {msg.get('message', '')}\n"
            history_text += f"Lumina: {msg.get('response', '')}\n\n"

        ai_response = get_ai_response(message, history_text)

        if not ai_response:
            ai_response = "Lumina could not generate a response."

        chat_collection.insert_one({
            "chat_id": chat_id,
            "user_id": user_id,
            "message": message,
            "response": ai_response,
            "created_at": now_iso()
        })

        try:
            chat_doc = conversations_collection.find_one({"_id": object_id(chat_id)})

            if chat_doc and chat_doc.get("title") == "New Chat":
                title = " ".join(message.split()[:6])

                conversations_collection.update_one(
                    {"_id": object_id(chat_id)},
                    {"$set": {"title": title}}
                )

        except Exception as e:
            print("TITLE UPDATE ERROR:", e)

        return {"response": ai_response}

    except HTTPException:
        raise

    except Exception as e:
        print("CHAT ROUTE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Chat failed")

@app.post("/chat-stream")
def chat_stream(data: ChatRequest):
    try:
        user_id = str(data.user_id).strip()
        chat_id = str(data.chat_id).strip()
        message = data.message.strip()

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        if not chat_id:
            raise HTTPException(status_code=400, detail="Chat ID is required")

        if not message:
            raise HTTPException(status_code=400, detail="Message is required")
        
        CHAT_MESSAGE_COST = CREDIT_COSTS["chat_message"]

        if not user_id.startswith("guest_"):
            deduct_user_credits(
            user_id,
            CHAT_MESSAGE_COST,
            "Chat message"
        )

        previous_messages = list(
            chat_collection
            .find({"chat_id": chat_id})
            .sort("_id", -1)
            .limit(6)
        )

        previous_messages.reverse()

        history_text = ""

        for msg in previous_messages:
            history_text += f"User: {msg.get('message', '')}\n"
            history_text += f"Lumina: {msg.get('response', '')}\n\n"

        def event_generator():
            full_response = ""

            try:
                for chunk in stream_ai_response(message, history_text):
                    full_response += chunk

                    yield f"data: {json.dumps({'chunk': chunk})}\n\n"

                if not full_response.strip():
                    full_response = "Lumina could not generate a response."

                chat_collection.insert_one({
                    "chat_id": chat_id,
                    "user_id": user_id,
                    "message": message,
                    "response": full_response,
                    "created_at": now_iso()
                })

                try:
                    chat_doc = conversations_collection.find_one({
                        "_id": object_id(chat_id)
                    })

                    if chat_doc and chat_doc.get("title") == "New Chat":
                        title = " ".join(message.split()[:6])

                        conversations_collection.update_one(
                            {"_id": object_id(chat_id)},
                            {"$set": {"title": title}}
                        )

                except Exception as e:
                    print("TITLE UPDATE ERROR:", e)

                yield f"data: {json.dumps({'done': True})}\n\n"

            except Exception as e:
                print("CHAT STREAM ERROR:", e)
                traceback.print_exc()

                yield f"data: {json.dumps({'error': str(e)})}\n\n"

        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream"
        )

    except HTTPException:
        raise

    except Exception as e:
        print("CHAT STREAM ROUTE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Chat stream failed")

@app.get("/chats/{user_id}")
def get_chats(user_id: str):
    try:
        chats = list(
            conversations_collection
            .find({
                "user_id": str(user_id),
                "archived": {"$ne": True}
            })
            .sort([
                ("pinned", -1),
                ("created_at", -1)
            ])
        )

        return serialize_docs(chats)

    except Exception as e:
        print("GET CHATS ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load chats")


@app.get("/messages/{chat_id}")
def get_messages(chat_id: str):
    try:
        messages = list(
            chat_collection
            .find({"chat_id": str(chat_id)})
            .sort("_id", 1)
        )

        return serialize_docs(messages)

    except Exception as e:
        print("GET MESSAGES ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load messages")


@app.put("/rename-chat/{chat_id}")
def rename_chat(chat_id: str, data: dict):
    try:
        new_title = data.get("title", "").strip()

        if not new_title:
            raise HTTPException(status_code=400, detail="Title is required")

        result = conversations_collection.update_one(
            {"_id": object_id(chat_id)},
            {"$set": {"title": new_title}}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Chat not found")

        return {
            "message": "Chat renamed successfully",
            "title": new_title
        }

    except HTTPException:
        raise

    except Exception as e:
        print("RENAME CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not rename chat")


@app.delete("/clear-chat/{chat_id}")
def clear_chat(chat_id: str):
    try:
        chat_collection.delete_many({"chat_id": str(chat_id)})

        return {"message": "Chat cleared successfully"}

    except Exception as e:
        print("CLEAR CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not clear chat")


@app.delete("/delete-chat/{chat_id}")
def delete_chat(chat_id: str):
    try:
        conversations_collection.delete_one({"_id": object_id(chat_id)})
        chat_collection.delete_many({"chat_id": str(chat_id)})

        return {"message": "Chat deleted"}

    except Exception as e:
        print("DELETE CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete chat")


@app.put("/pin-chat/{chat_id}")
def pin_chat(chat_id: str):
    try:
        chat = conversations_collection.find_one({"_id": object_id(chat_id)})

        if not chat:
            raise HTTPException(status_code=404, detail="Chat not found")

        new_status = not chat.get("pinned", False)

        conversations_collection.update_one(
            {"_id": object_id(chat_id)},
            {"$set": {"pinned": new_status}}
        )

        return {
            "message": "Pin status updated",
            "pinned": new_status
        }

    except HTTPException:
        raise

    except Exception as e:
        print("PIN CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not pin chat")


@app.put("/archive-chat/{chat_id}")
def archive_chat(chat_id: str):
    try:
        result = conversations_collection.update_one(
            {"_id": object_id(chat_id)},
            {"$set": {"archived": True}}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Chat not found")

        return {"message": "Chat archived"}

    except HTTPException:
        raise

    except Exception as e:
        print("ARCHIVE CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not archive chat")


@app.put("/unarchive-chat/{chat_id}")
def unarchive_chat(chat_id: str):
    try:
        result = conversations_collection.update_one(
            {"_id": object_id(chat_id)},
            {"$set": {"archived": False}}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Chat not found")

        return {"message": "Chat unarchived"}

    except HTTPException:
        raise

    except Exception as e:
        print("UNARCHIVE CHAT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not unarchive chat")


@app.get("/archived-chats/{user_id}")
def get_archived_chats(user_id: str):
    try:
        chats = list(
            conversations_collection
            .find({
                "user_id": str(user_id),
                "archived": True
            })
            .sort("created_at", -1)
        )

        return serialize_docs(chats)

    except Exception as e:
        print("ARCHIVED CHATS ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load archived chats")


@app.delete("/clear-all-chats/{user_id}")
def clear_all_chats(user_id: str):
    try:
        conversations_collection.delete_many({"user_id": str(user_id)})
        chat_collection.delete_many({"user_id": str(user_id)})

        return {"message": "All chats cleared"}

    except Exception as e:
        print("CLEAR ALL CHATS ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not clear chats")


# =========================
# CHAT WITH FILE
# =========================

def extract_text_from_uploaded_file(file_bytes: bytes, filename: str):
    filename = filename.lower()

    if filename.endswith(".pdf"):
        try:
            pdf_reader = PdfReader(BytesIO(file_bytes))
            extracted_text = ""

            for page in pdf_reader.pages:
                extracted_text += page.extract_text() or ""

            return extracted_text.strip()

        except Exception as e:
            print("PDF READ ERROR:", e)
            traceback.print_exc()
            return ""

    text_extensions = (
        ".txt", ".csv", ".json", ".py", ".js",
        ".html", ".css", ".md", ".xml"
    )

    if filename.endswith(text_extensions):
        try:
            return file_bytes.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            print("TEXT FILE READ ERROR:", e)
            traceback.print_exc()
            return ""

    return ""


@app.post("/chat-with-file")
async def chat_with_file(
    user_id: str = Form(...),
    chat_id: str = Form(...),
    message: str = Form("Analyze this file."),
    file: UploadFile = File(...)
):
    try:
        ensure_gemini_ready()

        user_id = str(user_id).strip()
        chat_id = str(chat_id).strip()
        message = (message or "").strip() or "Analyze this file."

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        if not chat_id:
            raise HTTPException(status_code=400, detail="Chat ID is required")

        file_bytes = await file.read()
        filename = file.filename or "uploaded_file"
        content_type = file.content_type or ""

        print("=== CHAT WITH FILE DEBUG ===")
        print("Filename:", filename)
        print("Content type:", content_type)
        print("Message:", message)
        print("File size:", len(file_bytes))

        if not file_bytes:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")

        if len(file_bytes) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Max size is 10MB.")

        ai_response = ""

        if content_type.startswith("image/"):
            prompt = f"""
You are Lumina AI.

The user uploaded an image.

User question:
{message}

Analyze the image carefully and answer clearly.
Mention important visible details, objects, colors, scene, and any visible text.
"""

            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=[
                    prompt,
                    types.Part.from_bytes(
                        data=file_bytes,
                        mime_type=content_type
                    )
                ],
                config=types.GenerateContentConfig(
                    temperature=0.4,
                    max_output_tokens=800
                )
            )

            ai_response = response.text.strip() if response and response.text else "I could not analyze this image properly."

        else:
            extracted_text = extract_text_from_uploaded_file(file_bytes, filename)

            if not extracted_text:
                raise HTTPException(
                    status_code=400,
                    detail="Could not read this file. Please upload a readable PDF, text file, or code file."
                )

            extracted_text = extracted_text[:12000]

            prompt = f"""
You are Lumina AI.

The user uploaded a file.

File name:
{filename}

User question:
{message}

File content:
{extracted_text}

Task:
Analyze this file and answer clearly.
If it is a document, summarize the important information.
If it is code, explain what it does and mention errors if present.
"""

            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.5,
                    max_output_tokens=900
                )
            )

            ai_response = response.text.strip() if response and response.text else "I could not analyze this file properly."

        chat_collection.insert_one({
            "chat_id": chat_id,
            "user_id": user_id,
            "message": f"{message}\n\nUploaded file: {filename}",
            "response": ai_response,
            "uploaded_file": filename,
            "created_at": now_iso()
        })

        try:
            chat_doc = conversations_collection.find_one({"_id": object_id(chat_id)})

            if chat_doc and chat_doc.get("title") == "New Chat":
                title = " ".join(message.split()[:6])
                conversations_collection.update_one(
                    {"_id": object_id(chat_id)},
                    {"$set": {"title": title}}
                )
        except Exception as e:
            print("FILE CHAT TITLE UPDATE ERROR:", e)

        return {"response": ai_response}

    except HTTPException:
        raise

    except Exception as e:
        print("CHAT WITH FILE GENERAL ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"File analysis failed: {str(e)}")


# =========================
# EXPORT CHAT
# =========================

@app.get("/export-chat-txt/{chat_id}")
def export_chat_txt(chat_id: str):
    try:
        messages = list(
            chat_collection
            .find({"chat_id": str(chat_id)})
            .sort("_id", 1)
        )

        if not messages:
            raise HTTPException(status_code=404, detail="No messages found")

        chat_text = "Lumina AI Chat Export\n"
        chat_text += "=" * 40 + "\n\n"

        for msg in messages:
            chat_text += f"User: {msg.get('message', '')}\n\n"
            chat_text += f"Lumina: {msg.get('response', '')}\n\n"
            chat_text += "-" * 40 + "\n\n"

        return Response(
            content=chat_text,
            media_type="text/plain",
            headers={
                "Content-Disposition": "attachment; filename=lumina_chat.txt"
            }
        )

    except HTTPException:
        raise

    except Exception as e:
        print("EXPORT TXT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not export chat")


@app.get("/export-chat-pdf/{chat_id}")
def export_chat_pdf(chat_id: str):
    try:
        messages = list(
            chat_collection
            .find({"chat_id": str(chat_id)})
            .sort("_id", 1)
        )

        if not messages:
            raise HTTPException(status_code=404, detail="No messages found")

        file_path = os.path.join(EXPORT_DIR, f"lumina_chat_{chat_id}.pdf")

        pdf = canvas.Canvas(file_path, pagesize=A4)
        width, height = A4

        x = 0.75 * inch
        y = height - 0.8 * inch

        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(x, y, "Lumina AI Chat Export")

        y -= 0.35 * inch

        pdf.setFont("Helvetica", 9)
        pdf.drawString(x, y, f"Chat ID: {chat_id}")

        y -= 0.4 * inch

        def write_wrapped_text(label, text, y_position):
            pdf.setFont("Helvetica-Bold", 11)
            pdf.drawString(x, y_position, label)

            y_position -= 0.22 * inch

            pdf.setFont("Helvetica", 10)

            clean_text = str(text).replace("\n", " ")

            wrapped_lines = textwrap.wrap(clean_text, width=90)

            for line in wrapped_lines:
                if y_position < 0.8 * inch:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    y_position = height - 0.8 * inch

                pdf.drawString(x, y_position, line)
                y_position -= 0.18 * inch

            return y_position - 0.15 * inch

        for msg in messages:
            if y < 1.2 * inch:
                pdf.showPage()
                y = height - 0.8 * inch

            y = write_wrapped_text("User:", msg.get("message", ""), y)
            y = write_wrapped_text("Lumina:", msg.get("response", ""), y)

            pdf.setStrokeColorRGB(0.6, 0.6, 0.6)
            pdf.line(x, y, width - x, y)
            y -= 0.25 * inch

        pdf.save()

        return FileResponse(
            path=file_path,
            media_type="application/pdf",
            filename="lumina_chat.pdf"
        )

    except HTTPException:
        raise

    except Exception as e:
        print("EXPORT PDF ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not export chat as PDF")


# =========================
# IMAGE GENERATION
# =========================

@app.post("/generate-image")
async def generate_image(
    prompt: str = Form(""),
    user_id: str | None = Form(None),
    style: str = Form("Realistic"),
    aspect_ratio: str = Form("Square"),
    image: UploadFile | None = File(None)
):
    try:
        prompt = (prompt or "").strip()

        if not prompt and image is None:
            raise HTTPException(
                status_code=400,
                detail="Please enter a prompt or upload a reference image."
            )

        if not GEMINI_API_KEY or not gemini_client:
            raise HTTPException(status_code=500, detail="GEMINI_API_KEY is missing in .env")

        ratio_map = {
            "Square": (1024, 1024),
            "Portrait": (1024, 1280),
            "Landscape": (1280, 720),
            "1:1": (1024, 1024),
            "4:5": (1024, 1280),
            "16:9": (1280, 720)
        }

        width, height = ratio_map.get(aspect_ratio, (1024, 1024))

        image_description = ""

        if image is not None:
            try:
                image_bytes = await image.read()

                if not image_bytes:
                    raise HTTPException(status_code=400, detail="Uploaded image is empty.")

                if len(image_bytes) > 10 * 1024 * 1024:
                    raise HTTPException(status_code=400, detail="Image is too large. Upload an image under 10 MB.")

                content_type = image.content_type or "image/jpeg"

                analysis_prompt = """
Analyze this uploaded image for reference-based AI image generation.

Return a strict visual description with these sections:

Subject:
Pose / composition:
Clothing / appearance:
Background:
Lighting:
Colors:
Camera angle:
Style:
Important details to preserve:
Things that must NOT change:

Be specific and concise.
"""

                analysis = gemini_client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=[
                        analysis_prompt,
                        types.Part.from_bytes(
                            data=image_bytes,
                            mime_type=content_type
                        )
                    ],
                    config=types.GenerateContentConfig(
                        temperature=0.4,
                        max_output_tokens=600
                    )
                )

                image_description = analysis.text.strip() if analysis and analysis.text else ""

            except HTTPException:
                raise

            except Exception as e:
                print("REFERENCE IMAGE ANALYSIS ERROR:", e)
                traceback.print_exc()
                image_description = ""

        style_map = {
            "Realistic": "photorealistic, realistic lighting, sharp details",
            "Anime": "anime style, clean line art, vibrant colors",
            "Cyberpunk": "cyberpunk, neon lights, futuristic atmosphere",
            "3D": "3D rendered, detailed materials, cinematic lighting",
            "Logo": "minimal logo design, clean vector style",
            "Watercolor": "watercolor painting, soft brush texture"
        }

        style_text = style_map.get(style, "high quality, detailed")

        if image_description:
            final_prompt = f"""
REFERENCE IMAGE ANALYSIS:
{image_description}

USER REQUEST:
{prompt or "Create a high-quality image using the uploaded image as reference."}

STYLE:
{style_text}

STRICT RULES:
- Preserve the main subject from the reference image.
- Preserve pose, composition, clothing, and background layout unless the user asks to change them.
- Do not add extra people or unrelated objects.
- Do not completely change the identity or main structure.
- Improve quality, lighting, detail, and style only.
- Avoid blurry, distorted face, extra fingers, extra limbs, watermark, text, logo, cropped face.
- Create a sharp, clean, high-quality image.
"""
        else:
            final_prompt = f"""
USER REQUEST:
{prompt}

STYLE:
{style_text}

RULES:
- Follow the user prompt closely.
- Make it high-quality, sharp, detailed, and visually attractive.
- Clean composition, professional lighting.
- Avoid blurry, distorted anatomy, watermark, text, logo.
"""

        encoded_prompt = urllib.parse.quote(final_prompt)

        image_api_url = (
            f"https://image.pollinations.ai/prompt/{encoded_prompt}"
            f"?width={width}&height={height}"
            f"&seed={int(time.time())}"
            f"&nologo=true"
        )

        response = requests.get(image_api_url, timeout=120)

        if response.status_code != 200:
            print("POLLINATIONS IMAGE ERROR:", response.status_code, response.text[:500])
            raise HTTPException(status_code=500, detail="Image generation failed.")

        filename = f"lumina_image_{uuid.uuid4().hex}.jpg"
        file_path = os.path.join(GENERATED_IMAGE_DIR, filename)

        with open(file_path, "wb") as f:
            f.write(response.content)

        image_url = f"/generated_images/{filename}"

        if user_id:
            try:
                images_collection.insert_one({
                    "user_id": str(user_id),
                    "prompt": prompt or "Generated from uploaded reference image",
                    "style": style,
                    "aspect_ratio": aspect_ratio,
                    "size": f"{width}x{height}",
                    "reference_used": bool(image_description),
                    "reference_description": image_description,
                    "image_url": image_url,
                    "created_at": now_iso()
                })
            except Exception as db_error:
                print("IMAGE HISTORY SAVE ERROR:", db_error)
                traceback.print_exc()

        return {
            "success": True,
            "message": "Image generated successfully",
            "image_url": image_url,
            "prompt_used": prompt,
            "style": style,
            "aspect_ratio": aspect_ratio
        }

    except HTTPException:
        raise

    except Exception as e:
        print("GENERATE IMAGE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Image generation failed: {str(e)}")


@app.get("/image-history/{user_id}")
def get_image_history(user_id: str):
    try:
        images = list(
            images_collection
            .find({"user_id": str(user_id)})
            .sort("_id", -1)
        )

        return serialize_docs(images)

    except Exception as e:
        print("IMAGE HISTORY ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load image history")


@app.delete("/delete-image/{image_id}")
def delete_image(image_id: str):
    try:
        image_doc = images_collection.find_one({"_id": object_id(image_id)})

        if not image_doc:
            raise HTTPException(status_code=404, detail="Image not found")

        image_url = image_doc.get("image_url", "")

        if image_url.startswith("/generated_images/"):
            file_name = image_url.replace("/generated_images/", "")
            file_path = os.path.join(GENERATED_IMAGE_DIR, file_name)

            if os.path.exists(file_path):
                os.remove(file_path)

        images_collection.delete_one({"_id": object_id(image_id)})

        return {"message": "Image deleted"}

    except HTTPException:
        raise

    except Exception as e:
        print("DELETE IMAGE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete image")


@app.delete("/delete-all-images/{user_id}")
def delete_all_images(user_id: str):
    try:
        images = list(images_collection.find({"user_id": str(user_id)}))

        for image in images:
            image_url = image.get("image_url", "")

            if image_url.startswith("/generated_images/"):
                file_name = image_url.replace("/generated_images/", "")
                file_path = os.path.join(GENERATED_IMAGE_DIR, file_name)

                if os.path.exists(file_path):
                    os.remove(file_path)

        images_collection.delete_many({"user_id": str(user_id)})

        return {"message": "All images deleted"}

    except Exception as e:
        print("DELETE ALL IMAGES ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete images")


# =========================
# VOICE STUDIO
# =========================

@app.post("/download-audio")
async def download_audio(data: dict):
    try:
        text = data.get("text")
        user_id = data.get("user_id", "")
        voice = data.get("voice", "en-US-AriaNeural")
        title = data.get("title", "Lumina Audio")

        if not text:
            raise HTTPException(status_code=400, detail="Text is required")

        filename = f"lumina_audio_{uuid.uuid4()}.mp3"
        file_path = os.path.join(GENERATED_AUDIO_DIR, filename)

        communicate = edge_tts.Communicate(text=text, voice=voice)

        audio_bytes = bytearray()

        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_bytes.extend(chunk["data"])

        if not audio_bytes:
            raise HTTPException(status_code=500, detail="No audio generated")

        with open(file_path, "wb") as f:
            f.write(bytes(audio_bytes))

        audio_url = f"/generated_audio/{filename}"

        audio_collection.insert_one({
            "user_id": str(user_id),
            "title": title,
            "text": text,
            "voice": voice,
            "audio_url": audio_url,
            "created_at": now_iso()
        })

        return FileResponse(
            path=file_path,
            media_type="audio/mpeg",
            filename=f"{title}.mp3"
        )

    except HTTPException:
        raise

    except Exception as e:
        print("AUDIO DOWNLOAD ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Audio generation failed")


@app.get("/audio-history/{user_id}")
def get_audio_history(user_id: str):
    try:
        audios = list(
            audio_collection
            .find({"user_id": str(user_id)})
            .sort("_id", -1)
        )

        return serialize_docs(audios)

    except Exception as e:
        print("AUDIO HISTORY ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load audio history")


@app.put("/rename-audio/{audio_id}")
def rename_audio(audio_id: str, data: dict):
    try:
        new_title = data.get("title", "").strip()

        if not new_title:
            raise HTTPException(status_code=400, detail="Title is required")

        result = audio_collection.update_one(
            {"_id": object_id(audio_id)},
            {"$set": {"title": new_title}}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Audio not found")

        return {
            "message": "Audio renamed successfully",
            "title": new_title
        }

    except HTTPException:
        raise

    except Exception as e:
        print("RENAME AUDIO ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not rename audio")


@app.delete("/delete-audio/{audio_id}")
def delete_audio(audio_id: str):
    try:
        audio_doc = audio_collection.find_one({"_id": object_id(audio_id)})

        if not audio_doc:
            raise HTTPException(status_code=404, detail="Audio not found")

        audio_url = audio_doc.get("audio_url", "")

        if audio_url.startswith("/generated_audio/"):
            file_name = audio_url.replace("/generated_audio/", "")
            file_path = os.path.join(GENERATED_AUDIO_DIR, file_name)

            if os.path.exists(file_path):
                os.remove(file_path)

        audio_collection.delete_one({"_id": object_id(audio_id)})

        return {"message": "Audio deleted"}

    except HTTPException:
        raise

    except Exception as e:
        print("DELETE AUDIO ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete audio")


@app.delete("/delete-all-audio/{user_id}")
def delete_all_audio(user_id: str):
    try:
        audios = list(audio_collection.find({"user_id": str(user_id)}))

        for audio in audios:
            audio_url = audio.get("audio_url", "")

            if audio_url.startswith("/generated_audio/"):
                file_name = audio_url.replace("/generated_audio/", "")
                file_path = os.path.join(GENERATED_AUDIO_DIR, file_name)

                if os.path.exists(file_path):
                    os.remove(file_path)

        audio_collection.delete_many({"user_id": str(user_id)})

        return {"message": "All audio history deleted"}

    except Exception as e:
        print("DELETE ALL AUDIO ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete audio history")

# =========================
# AI VIDEO STUDIO
# =========================

def download_file_from_url(url: str, save_path: str):
    try:
        response = requests.get(url, timeout=180)

        if response.status_code != 200:
            print("VIDEO DOWNLOAD ERROR:", response.status_code, response.text[:300])
            raise HTTPException(
                status_code=500,
                detail="Could not download generated video."
            )

        with open(save_path, "wb") as f:
            f.write(response.content)

        return save_path

    except HTTPException:
        raise

    except Exception as e:
        print("DOWNLOAD VIDEO FILE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="Video download failed."
        )


# =========================
# NOTES / SAVED RESPONSES
# =========================

@app.post("/save-note")
def save_note(data: dict):
    try:
        user_id = str(data.get("user_id", "")).strip()
        title = str(data.get("title", "")).strip()
        content = str(data.get("content", "")).strip()
        category = str(data.get("category", "General")).strip()

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        if not content:
            raise HTTPException(status_code=400, detail="Note content is required")

        if not title:
            title = " ".join(content.split()[:8]) or "Saved Note"

        note = {
            "user_id": user_id,
            "title": title[:80],
            "content": content,
            "category": category or "General",
            "created_at": now_iso()
        }

        result = notes_collection.insert_one(note)

        return {
            "message": "Note saved successfully",
            "note_id": str(result.inserted_id)
        }

    except HTTPException:
        raise

    except Exception as e:
        print("SAVE NOTE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not save note")


@app.get("/notes/{user_id}")
def get_notes(user_id: str):
    try:
        notes = list(
            notes_collection
            .find({"user_id": str(user_id)})
            .sort("_id", -1)
        )

        return serialize_docs(notes)

    except Exception as e:
        print("GET NOTES ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load notes")


@app.delete("/delete-note/{note_id}")
def delete_note(note_id: str):
    try:
        result = notes_collection.delete_one({"_id": object_id(note_id)})

        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Note not found")

        return {"message": "Note deleted"}

    except HTTPException:
        raise

    except Exception as e:
        print("DELETE NOTE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete note")


@app.delete("/delete-all-notes/{user_id}")
def delete_all_notes(user_id: str):
    try:
        notes_collection.delete_many({"user_id": str(user_id)})

        return {"message": "All notes deleted"}

    except Exception as e:
        print("DELETE ALL NOTES ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not delete notes")


# =========================
# DASHBOARD
# =========================

@app.get("/dashboard-stats/{user_id}")
def dashboard_stats(user_id: str):
    try:
        user_id = str(user_id)

        chats_count = conversations_collection.count_documents({"user_id": user_id})
        images_count = images_collection.count_documents({"user_id": user_id})
        audio_count = audio_collection.count_documents({"user_id": user_id})

        latest_chat = chat_collection.find_one(
            {"user_id": user_id},
            sort=[("_id", -1)]
        )

        latest_image = images_collection.find_one(
            {"user_id": user_id},
            sort=[("_id", -1)]
        )

        latest_audio = audio_collection.find_one(
            {"user_id": user_id},
            sort=[("_id", -1)]
        )

        latest_dates = []

        if latest_chat and latest_chat.get("created_at"):
            latest_dates.append(latest_chat.get("created_at"))

        if latest_image and latest_image.get("created_at"):
            latest_dates.append(latest_image.get("created_at"))

        if latest_audio and latest_audio.get("created_at"):
            latest_dates.append(latest_audio.get("created_at"))

        last_active = "No activity yet"

        if latest_dates:
            last_active = max(latest_dates)

        return {
            "chats_created": chats_count,
            "images_generated": images_count,
            "audio_created": audio_count,
            "last_active": last_active
        }

    except Exception as e:
        print("DASHBOARD STATS ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load dashboard stats")


# =========================
# GLOBAL SEARCH
# =========================

@app.get("/global-search/{user_id}")
def global_search(user_id: str, q: str = ""):
    try:
        user_id = str(user_id)
        query = q.strip()

        if not query:
            return {
                "chats": [],
                "messages": [],
                "images": [],
                "audio": [],
                "archives": []
            }

        search_regex = {
            "$regex": query,
            "$options": "i"
        }

        chats = list(
            conversations_collection
            .find({
                "user_id": user_id,
                "archived": {"$ne": True},
                "title": search_regex
            })
            .limit(10)
        )

        messages = list(
            chat_collection
            .find({
                "user_id": user_id,
                "$or": [
                    {"message": search_regex},
                    {"response": search_regex}
                ]
            })
            .limit(10)
        )

        images = list(
            images_collection
            .find({
                "user_id": user_id,
                "prompt": search_regex
            })
            .limit(10)
        )

        audio = list(
            audio_collection
            .find({
                "user_id": user_id,
                "$or": [
                    {"title": search_regex},
                    {"text": search_regex}
                ]
            })
            .limit(10)
        )

        archives = list(
            conversations_collection
            .find({
                "user_id": user_id,
                "archived": True,
                "title": search_regex
            })
            .limit(10)
        )

        return {
            "chats": serialize_docs(chats),
            "messages": serialize_docs(messages),
            "images": serialize_docs(images),
            "audio": serialize_docs(audio),
            "archives": serialize_docs(archives)
        }

    except Exception as e:
        print("GLOBAL SEARCH ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Global search failed")
    

# =========================
# DOCUMENT STUDIO
# =========================

def extract_text_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)

        return "\n".join(text_parts).strip()

    except Exception as e:
        print("PDF TEXT EXTRACTION ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=400,
            detail="Could not read PDF file."
        )


def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        text_parts = []

        # 1. Extract normal paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # 2. Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    text_parts.append(" | ".join(row_text))

        # 3. Extract text from headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

        # 4. Extract text from footers
        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

        final_text = "\n".join(text_parts).strip()

        print("DOCX EXTRACTED TEXT LENGTH:", len(final_text))
        print("DOCX EXTRACTED TEXT PREVIEW:", final_text[:500])

        return final_text

    except Exception as e:
        print("DOCX TEXT EXTRACTION ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=400,
            detail="Could not read DOCX file."
        )

def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

    except Exception as e:
        print("TXT TEXT EXTRACTION ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=400,
            detail="Could not read TXT file."
        )


def extract_document_text(file_path: str, filename: str) -> str:
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_path)

    if lower_name.endswith(".docx"):
        return extract_text_from_docx(file_path)

    if lower_name.endswith(".txt"):
        return extract_text_from_txt(file_path)

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Upload PDF, DOCX, or TXT."
    )


def build_document_prompt(document_text: str, task: str) -> str:
    short_text = document_text[:6000]

    task_instructions = {
        "summary": "Summarize this document clearly in structured headings and bullet points.",
        "simple": "Explain this document in very simple student-friendly language.",
        "key_points": "Extract the most important key points from this document.",
        "questions": "Create important exam-style questions and answers from this document.",
        "table": "Convert the important information into clean tables wherever possible."
    }

    instruction = task_instructions.get(
        task,
        "Analyze this document and explain it clearly."
    )

    return f"""
You are Lumina AI Document Studio.

Task:
{instruction}

Formatting rules:
- Use clear headings.
- Use bullet points where useful.
- Use markdown tables if comparison or structured data is present.
- Keep the answer organized and easy to study.
- Do not mention that the text was truncated.

Document text:
{short_text}
"""


@app.post("/upload-document")
async def upload_document(
    user_id: str = Form(...),
    task: str = Form("summary"),
    file: UploadFile = File(...)
):
    try:
        user_id = str(user_id).strip()
        analysis = None

        if not user_id:
            raise HTTPException(
                status_code=400,
                detail="User ID is required."
            )

        if user_id.startswith("guest_"):
            raise HTTPException(
                status_code=403,
                detail="Please login to use Document Studio."
            )

        if not file.filename:
            raise HTTPException(
                status_code=400,
                detail="File is required."
            )

        allowed_extensions = [".pdf", ".docx", ".txt"]

        if not any(file.filename.lower().endswith(ext) for ext in allowed_extensions):
            raise HTTPException(
                status_code=400,
                detail="Only PDF, DOCX, and TXT files are supported."
            )

        file_ext = os.path.splitext(file.filename)[1].lower()
        safe_filename = f"document_{uuid.uuid4().hex}{file_ext}"

        # If you moved uploads outside project, keep this:
        file_path = UPLOAD_DOCUMENTS_DIR / safe_filename

        file_bytes = await file.read()

        if not file_bytes:
            raise HTTPException(
                status_code=400,
                detail="Uploaded file is empty."
            )

        with open(file_path, "wb") as f:
            f.write(file_bytes)

        document_text = extract_document_text(file_path, file.filename)

        print("DOCUMENT TEXT LENGTH:", len(document_text))
        print("DOCUMENT TEXT PREVIEW:", document_text[:300])

        if not document_text or not document_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract readable text from this document."
            )

        prompt = build_document_prompt(document_text, task)

        try:
            analysis = get_ai_response(prompt, "")
        except Exception as ai_error:
            print("DOCUMENT AI ERROR:", ai_error)
            traceback.print_exc()
            analysis = None

        if not analysis:
            analysis = f"""
# Document Read Successfully

Lumina successfully uploaded and scanned your document, but the AI model could not generate a full analysis right now.

## File Name
{file.filename}

## Task
{task}

## Extracted Text Preview

{document_text[:1500]}

---

Please try again after a few seconds.
"""

        print("ANALYSIS RESULT TYPE:", type(analysis))
        print("ANALYSIS RESULT PREVIEW:", analysis[:300] if analysis else "NO ANALYSIS")

        doc_record = {
            "user_id": user_id,
            "original_filename": file.filename,
            "stored_filename": safe_filename,
            "file_url": f"/uploaded_documents/{safe_filename}",
            "task": task,
            "extracted_text_preview": document_text[:1000],
            "analysis": analysis,
            "created_at": now_iso()
        }

        result = documents_collection.insert_one(doc_record)

        return {
            "success": True,
            "document_id": str(result.inserted_id),
            "filename": file.filename,
            "task": task,
            "analysis": analysis,
            "file_url": f"/uploaded_documents/{safe_filename}"
        }

    except HTTPException:
        raise

    except Exception as e:
        print("UPLOAD DOCUMENT ERROR:", e)
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Document analysis failed: {str(e)}"
        )
    
@app.get("/document-history/{user_id}")
def get_document_history(user_id: str):
    try:
        docs = list(
            documents_collection
            .find({"user_id": str(user_id)})
            .sort("_id", -1)
        )

        return serialize_docs(docs)

    except Exception as e:
        print("DOCUMENT HISTORY ERROR:", e)
        raise HTTPException(
            status_code=500,
            detail="Could not load document history."
        )


@app.delete("/delete-document/{document_id}")
def delete_document(document_id: str):
    try:
        doc = documents_collection.find_one({
            "_id": object_id(document_id)
        })

        if not doc:
            raise HTTPException(
                status_code=404,
                detail="Document not found."
            )

        stored_filename = doc.get("stored_filename")

        if stored_filename:
            file_path = UPLOAD_DOCUMENTS_DIR / stored_filename

            if file_path.exists():
                file_path.unlink()

        documents_collection.delete_one({
            "_id": object_id(document_id)
        })

        return {
            "message": "Document deleted successfully."
        }

    except HTTPException:
        raise

    except Exception as e:
        print("DELETE DOCUMENT ERROR:", e)
        raise HTTPException(
            status_code=500,
            detail="Could not delete document."
        )


@app.delete("/delete-all-documents/{user_id}")
def delete_all_documents(user_id: str):
    try:
        docs = list(
            documents_collection.find({
                "user_id": str(user_id)
            })
        )

        for doc in docs:
            stored_filename = doc.get("stored_filename")

            if stored_filename:
                file_path = UPLOAD_DOCUMENTS_DIR / stored_filename

            if file_path.exists():
                file_path.unlink()

        documents_collection.delete_many({
            "user_id": str(user_id)
        })

        return {
            "message": "All documents deleted successfully."
        }

    except Exception as e:
        print("DELETE ALL DOCUMENTS ERROR:", e)
        raise HTTPException(
            status_code=500,
            detail="Could not delete documents."
        )

@app.get("/user-plan/{user_id}")
def get_user_plan(user_id: str):
    try:
        if str(user_id).startswith("guest_"):
            return {
                "success": True,
                "plan": "guest",
                "credits": 0,
                "monthly_credit_limit": 0
            }

        reset_monthly_credits_if_needed(user_id)

        user = users_collection.find_one({"_id": object_id(user_id)})

        if not user:
            raise HTTPException(status_code=404, detail="User not found.")

        return {
            "success": True,
            "plan": user.get("plan", "free"),
            "credits": int(user.get("credits", 0)),
            "monthly_credit_limit": int(user.get("monthly_credit_limit", 3000)),
            "daily_used_credits": int(user.get("daily_used_credits", 0)),
            "daily_credit_limit": int(user.get("daily_credit_limit", 300))
}

    except HTTPException:
        raise

    except Exception as e:
        print("GET USER PLAN ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not load plan.")
    
@app.post("/demo-upgrade-plan")
def demo_upgrade_plan(data: dict):
    try:
        user_id = str(data.get("user_id", "")).strip()
        plan = str(data.get("plan", "free")).lower().strip()

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required.")

        if user_id.startswith("guest_"):
            raise HTTPException(status_code=403, detail="Please login first.")

        if plan not in LUMINA_PLANS:
            raise HTTPException(status_code=400, detail="Invalid plan.")

        user = users_collection.find_one({"_id": object_id(user_id)})
        print("===== PLAN UPGRADE DEBUG =====")
        print("REQUEST USER ID:", user_id)
        print("REQUEST PLAN:", plan)
        print("DB USER EMAIL:", user.get("email") if user else None)
        print("DB EMAIL VERIFIED:", user.get("email_verified") if user else None)
        print("==============================")

        if not user:
            raise HTTPException(status_code=404, detail="User not found.")

        if plan != "free" and not user.get("email_verified", False):
            raise HTTPException(
                status_code=403,
                detail="Please verify your Gmail before purchasing a plan."
            )

        plan_details = LUMINA_PLANS[plan]

        update_data = {
            "plan": plan,
            "credits": plan_details["monthly_credits"],
            "monthly_credit_limit": plan_details["monthly_credits"],
            "daily_credit_limit": plan_details["daily_limit"],
            "daily_used_credits": 0,
            "last_credit_reset_month": current_month_key(),
            "last_daily_reset_date": current_date_key()
        }

        result = users_collection.update_one(
            {"_id": object_id(user_id)},
            {"$set": update_data}
        )

        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="User not found.")

        return {
            "success": True,
            "message": f"Plan upgraded to {plan_details['name']}.",
            "plan": plan,
            "credits": plan_details["monthly_credits"],
            "monthly_credit_limit": plan_details["monthly_credits"],
            "daily_credit_limit": plan_details["daily_limit"],
            "email_verified": user.get("email_verified", False)
        }

    except HTTPException:
        raise

    except Exception as e:
        print("DEMO UPGRADE ERROR:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Could not upgrade plan.")
    
    
@app.post("/admin/add-email-verification-fields")
def add_email_verification_fields():
    result = users_collection.update_many(
        {"email_verified": {"$exists": False}},
        {
            "$set": {
                "email_verified": False,
                "email_verification_token": "",
                "email_verification_expires": None
            }
        }
    )

    return {
        "message": "Email verification fields added.",
        "modified_count": result.modified_count
    }